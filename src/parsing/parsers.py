import json
import csv
import xml.etree.ElementTree as ET
from pathlib import Path
import sys
import os
import re
from datetime import datetime

import pdfplumber
from docx import Document
from openpyxl import load_workbook
import chardet
from datetime import datetime, timezone

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from storage.mongo import save_to_mongo


def normalize_text(text):
    if not text:
        return ""

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text)
    return text.strip()


def read_file_with_encoding(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()

    result = chardet.detect(raw)
    encoding = result.get("encoding") or "utf-8"
    confidence = result.get("confidence", 0)

    print(f"Detected encoding for {file_path}: {encoding} (confidence: {confidence})")

    try:
        text = raw.decode(encoding, errors="replace")
    except Exception:
        text = raw.decode("utf-8", errors="replace")

    return text


def parse_json_files():
    json_dir = Path("data/raw/api")
    parsed_articles = []

    if not json_dir.exists():
        print(f"JSON directory not found: {json_dir}")
        return parsed_articles

    json_files = sorted(json_dir.glob("*.json"))

    if not json_files:
        print("No JSON files found.")
        return parsed_articles

    for file_path in json_files:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        articles = data.get("articles", [])

        for article in articles:
            parsed_article = {
                "source": article.get("source", {}).get("name"),
                "author": article.get("author"),
                "title": article.get("title"),
                "description": article.get("description"),
                "url": article.get("url"),
                "publishedAt": article.get("publishedAt"),
                "document_type": "json",
                "file_name": file_path.name,
                "extraction_timestamp": datetime.now(timezone.utc).isoformat()
            }

            parsed_articles.append(parsed_article)
            save_to_mongo(parsed_article, file_path.name)

        print(f"Parsed {len(articles)} articles from {file_path.name}")

    print(f"Total parsed JSON articles: {len(parsed_articles)}")
    return parsed_articles


def parse_csv_file(file_path):
    parsed_rows = []

    with open(file_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row["document_type"] = "csv"
            row["file_name"] = Path(file_path).name
            row["extraction_timestamp"] = datetime.now(timezone.utc).isoformat()

            parsed_rows.append(row)
            save_to_mongo(row, Path(file_path).name)

    print(f"Parsed {len(parsed_rows)} rows from CSV: {file_path}")
    return parsed_rows


def parse_xml_file(file_path):
    parsed_items = []

    tree = ET.parse(file_path)
    root = tree.getroot()

    for article in root.findall("article"):
        parsed_article = {
            "id": article.findtext("id"),
            "title": article.findtext("title"),
            "category": article.findtext("category"),
            "document_type": "xml",
            "file_name": Path(file_path).name,
            "extraction_timestamp": datetime.now(timezone.utc).isoformat()
        }
        parsed_items.append(parsed_article)
        save_to_mongo(parsed_article, Path(file_path).name)

    print(f"Parsed {len(parsed_items)} items from XML: {file_path}")
    return parsed_items


def extract_text_from_pdf(pdf_path):
    extracted_pages = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            tables = page.extract_tables()

            if text or tables:
                page_data = {
                    "file_name": Path(pdf_path).name,
                    "document_type": "pdf",
                    "page_number": page_num,
                    "text": normalize_text(text or ""),
                    "tables": tables if tables else [],
                    "source": str(pdf_path),
                    "extraction_timestamp": datetime.utcnow().isoformat(),
                    "extraction_library": "pdfplumber"
                }
                extracted_pages.append(page_data)

    print(f"Extracted {len(extracted_pages)} pages from PDF: {pdf_path}")
    return extracted_pages


def extract_text_from_two_column_pdf(pdf_path, gap=10):
    extracted_pages = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            mid_x = page.width / 2

            left_column = page.crop((0, 0, mid_x - gap, page.height))
            right_column = page.crop((mid_x + gap, 0, page.width, page.height))

            left_text = normalize_text(left_column.extract_text() or "")
            right_text = normalize_text(right_column.extract_text() or "")
            combined_text = "\n\n".join(part for part in [left_text, right_text] if part)

            tables = page.extract_tables()

            if combined_text or tables:
                page_data = {
                    "file_name": Path(pdf_path).name,
                    "document_type": "pdf_two_column",
                    "page_number": page_num,
                    "text": combined_text,
                    "tables": tables if tables else [],
                    "source": str(pdf_path),
                    "extraction_timestamp": datetime.utcnow().isoformat(),
                    "extraction_library": "pdfplumber"
                }
                extracted_pages.append(page_data)

    print(f"Extracted {len(extracted_pages)} pages from two-column PDF: {pdf_path}")
    return extracted_pages


def extract_text_from_word(docx_path):
    doc = Document(docx_path)

    paragraphs = []
    tables = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    extracted_data = {
        "file_name": Path(docx_path).name,
        "document_type": "word",
        "text": "\n\n".join(paragraphs),
        "tables": tables,
        "source": str(docx_path),
        "extraction_timestamp": datetime.utcnow().isoformat(),
        "extraction_library": "python-docx"
    }

    print(f"Extracted text from Word file: {docx_path}")
    return extracted_data


def extract_text_from_two_column_word(docx_path):
    doc = Document(docx_path)

    paragraphs = []
    tables = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    extracted_data = {
        "file_name": Path(docx_path).name,
        "document_type": "word_two_column",
        "text": "\n\n".join(paragraphs),
        "tables": tables,
        "source": str(docx_path),
        "extraction_timestamp": datetime.utcnow().isoformat(),
        "extraction_library": "python-docx"
    }

    print(f"Extracted text from two-column Word file: {docx_path}")
    return extracted_data


def extract_word_runs(docx_path):
    doc = Document(docx_path)
    runs_data = []

    for para_index, para in enumerate(doc.paragraphs, start=1):
        for run_index, run in enumerate(para.runs, start=1):
            if run.text.strip():
                runs_data.append({
                    "paragraph_number": para_index,
                    "run_number": run_index,
                    "text": run.text.strip(),
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline
                })

    print(f"Extracted {len(runs_data)} runs from Word file: {docx_path}")
    return runs_data


def extract_data_from_excel(file_path):
    wb = load_workbook(file_path)
    ws = wb["Articles"]

    articles = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        if all(cell is None for cell in row):
            continue

        article = {
            "id": row[0],
            "title": row[1],
            "source": row[2],
            "category": row[3],
            "published_date": row[4],
            "mentions": row[5],
            "sentiment_score": row[6],
            "document_type": "excel",
            "sheet_name": "Articles",
            "file_name": Path(file_path).name,
            "extraction_timestamp": datetime.utcnow().isoformat(),
            "extraction_library": "openpyxl"
        }
        articles.append(article)

    print(f"Extracted {len(articles)} rows from Excel sheet 'Articles'")
    return articles


def extract_summary_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb["Summary"]

    summary = {
        "document_type": "excel_summary",
        "sheet_name": "Summary",
        "file_name": Path(file_path).name,
        "extraction_timestamp": datetime.utcnow().isoformat(),
        "extraction_library": "openpyxl"
    }

    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is not None:
            summary[row[0]] = row[1]

    print(f"Extracted summary from Excel sheet 'Summary'")
    return summary


if __name__ == "__main__":
    # Existing API / structured data parsing
    json_data = parse_json_files()
    print("First 2 parsed JSON articles:")
    print(json_data[:2])

    # PDF
    normal_pdf = "data/raw/pdf/news_normal.pdf"
    two_column_pdf = "data/raw/pdf/news_two_column.pdf"

    pdf_data = extract_text_from_pdf(normal_pdf)
    print("First normal PDF page:")
    print(pdf_data[:1])

    pdf_two_col_data = extract_text_from_two_column_pdf(two_column_pdf)
    print("First two-column PDF page:")
    print(pdf_two_col_data[:1])

    for page in pdf_data:
        save_to_mongo(page, page["file_name"])

    for page in pdf_two_col_data:
        save_to_mongo(page, page["file_name"])

    # Word
    normal_word = "data/raw/word/news_normal.docx"
    two_column_word = "data/raw/word/news_two_column.docx"

    word_data = extract_text_from_word(normal_word)
    print("Normal Word extraction:")
    print(word_data)

    word_two_col_data = extract_text_from_two_column_word(two_column_word)
    print("Two-column Word extraction:")
    print(word_two_col_data)

    word_runs = extract_word_runs(normal_word)
    print("First 5 Word runs:")
    print(word_runs[:5])

    save_to_mongo(word_data, word_data["file_name"])
    save_to_mongo(word_two_col_data, word_two_col_data["file_name"])

    # Excel
    excel_path = "data/raw/excel/news_data.xlsx"

    excel_articles = extract_data_from_excel(excel_path)
    print("Excel Articles:")
    print(excel_articles)

    excel_summary = extract_summary_from_excel(excel_path)
    print("Excel Summary:")
    print(excel_summary)

    for article in excel_articles:
        save_to_mongo(article, "news_data.xlsx")

    save_to_mongo(excel_summary, "news_data.xlsx")

    # Encoding test
    encoding_text = read_file_with_encoding("data/raw/api/news_page_1.json")
    print("Encoding test:")
    print(encoding_text[:200])